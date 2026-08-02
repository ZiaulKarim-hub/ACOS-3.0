#!/usr/bin/env python3
"""
acos-skill-breakdown — structured breakdown JSON  ->  styled Microsoft Word (.docx)

WHY PYTHON (per repo rule: new code defaults to TypeScript/Rust; Python only when
unavoidable). Generating a real .docx with controlled fonts, table shading and
column widths requires python-docx — the Word document format has no viable
TS/Rust library of equivalent fidelity, and the project's own DOCX-quality
preference names python-docx as the required path. This file does one thing:
render an already-decided breakdown into Word. All judgement (which steps exist,
the 5-word labels, the plain-language explanations) is made upstream by the model
and handed in as JSON — this renderer never invents content.

INPUT  (JSON, from --in <file> or stdin) — schema:
{
  "skill":     "acos-research-riffs",            # required, the skill's name
  "title":     "…",                              # optional; defaults from skill
  "subtitle":  "…",                              # optional one-line what-it-is
  "source":    ".claude/skills/…/SKILL.md",      # optional provenance path
  "generated": "2026-07-23",                      # optional ISO date string
  "phases": [                                     # required, >=1
    {
      "id":   "P0",                               # optional short tag
      "name": "Phase 0 — Preflight",              # required
      "purpose": "one line: why this phase runs", # optional
      "steps": [                                  # required, >=1
        {"id":"0.1","label":"exactly five words here","explanation":"plain text"},
        ...
      ]
    },
    ...
  ]
}

USAGE:
  python3 render_docx.py --in breakdown.json --out /path/Skill-Breakdown.docx
  cat breakdown.json | python3 render_docx.py --out out.docx

The 5-word rule is checked, not enforced destructively: any label whose word
count != 5 is listed on stderr and marked in the document, so a slip is visible
rather than silent. Exit code is 0 on success even with warnings; non-zero only
on a schema/IO error, so a caller can branch on real failures.
"""

import argparse
import json
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_SECTION
    from docx.oxml.ns import nsdecls, qn
    from docx.oxml import parse_xml, OxmlElement
except ImportError:
    sys.stderr.write(
        "ERROR: python-docx is required. Install with: pip install python-docx\n"
    )
    sys.exit(2)


# ── OKOA institutional palette (matches .claude/scripts/html-to-docx.py) ──────
SAGE_80   = RGBColor(0x46, 0x5D, 0x53)   # primary header / heading ink
SAGE_60   = RGBColor(0x6E, 0x8B, 0x7F)   # muted secondary
INK       = RGBColor(0x16, 0x16, 0x16)   # body text
CORAL     = RGBColor(0xFF, 0x79, 0x5E)   # accent
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BG = "465D53"                      # sage-80  (table header row)
ZEBRA_BG  = "EEF3F0"                      # sage-05  (alternating rows)
PHASE_BG  = "DCE7E1"                      # sage-10  (phase strip)
RULE      = "C9D8D0"                      # hairline border

HEADING_FONT = "IBM Plex Sans"
BODY_FONT    = "IBM Plex Sans"


def _shade(cell, color_hex):
    """Fill a table cell background with a solid hex color."""
    tcPr = cell._tc.get_or_add_tcPr()
    for old in tcPr.findall(qn("w:shd")):
        tcPr.remove(old)
    tcPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'))


def _set_borders(cell, color_hex=RULE, sz=4):
    """Give a cell thin borders on all four sides."""
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color_hex)
        borders.append(e)
    tcPr.append(borders)


def _col_widths(table, widths_in):
    """Force fixed column widths (python-docx needs it set per-cell)."""
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_in):
            cell.width = Inches(w)


def _cell_text(cell, text, *, bold=False, color=INK, size=10.5, font=BODY_FONT,
               align=WD_ALIGN_PARAGRAPH.LEFT, space_after=2):
    """Write text into a cell as a single styled run, replacing any default."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = font
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p


def _word_count(label: str) -> int:
    return len(re.findall(r"\S+", label.strip()))


def _default_style(doc):
    st = doc.styles["Normal"]
    st.font.name = BODY_FONT
    st.font.size = Pt(10.5)
    st.font.color.rgb = INK


def _add_footer(section, skill, generated):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.text = ""
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    bits = [f"/{skill} — step-by-step breakdown"]
    if generated:
        bits.append(f"generated {generated}")
    run = p.add_run("     ·     ".join(bits))
    run.font.name = BODY_FONT
    run.font.size = Pt(8)
    run.font.color.rgb = SAGE_60


def render(data: dict, out_path: Path) -> dict:
    skill = data.get("skill") or "unknown-skill"
    title = data.get("title") or f"/{skill} — Step-by-Step Breakdown"
    subtitle = data.get("subtitle") or ""
    source = data.get("source") or ""
    generated = data.get("generated") or ""
    phases = data.get("phases") or []
    if not phases:
        raise ValueError("breakdown JSON has no phases")

    doc = Document()
    _default_style(doc)

    sec = doc.sections[0]
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.85)
    sec.right_margin = Inches(0.85)
    usable = sec.page_width.inches - sec.left_margin.inches - sec.right_margin.inches
    _add_footer(sec, skill, generated)

    # ── Title block ──────────────────────────────────────────────────────────
    h = doc.add_paragraph()
    h.paragraph_format.space_after = Pt(2)
    r = h.add_run(title)
    r.font.name = HEADING_FONT
    r.font.size = Pt(20)
    r.bold = True
    r.font.color.rgb = SAGE_80

    if subtitle:
        s = doc.add_paragraph()
        s.paragraph_format.space_after = Pt(6)
        rs = s.add_run(subtitle)
        rs.font.name = BODY_FONT
        rs.font.size = Pt(11)
        rs.italic = True
        rs.font.color.rgb = SAGE_60

    meta_bits = []
    if source:
        meta_bits.append(f"Source: {source}")
    if generated:
        meta_bits.append(f"Generated: {generated}")
    total_steps = sum(len(p.get("steps") or []) for p in phases)
    meta_bits.append(f"{len(phases)} phases · {total_steps} steps")
    m = doc.add_paragraph()
    m.paragraph_format.space_after = Pt(10)
    rm = m.add_run("      ".join(meta_bits))
    rm.font.name = BODY_FONT
    rm.font.size = Pt(8.5)
    rm.font.color.rgb = SAGE_60

    # ── How to read ──────────────────────────────────────────────────────────
    legend = doc.add_paragraph()
    legend.paragraph_format.space_after = Pt(12)
    lr = legend.add_run(
        "How to read this: each row is one atomic step of the skill, taken in order. "
        "The middle column names the step in exactly five words; the right column "
        "explains, in plain language, what that step does and why it matters."
    )
    lr.font.name = BODY_FONT
    lr.font.size = Pt(9.5)
    lr.font.color.rgb = INK

    # Column widths: #  |  Step (5 words)  |  What it means
    col_w = [0.55, round(usable * 0.32, 2), round(usable * 0.58, 2)]
    col_w[2] = round(usable - col_w[0] - col_w[1], 2)

    warnings = []

    for phase in phases:
        pname = phase.get("name") or phase.get("id") or "Phase"
        purpose = phase.get("purpose") or ""
        steps = phase.get("steps") or []

        # Phase heading strip
        ph = doc.add_paragraph()
        ph.paragraph_format.space_before = Pt(10)
        ph.paragraph_format.space_after = Pt(1)
        pr = ph.add_run(pname)
        pr.font.name = HEADING_FONT
        pr.font.size = Pt(13)
        pr.bold = True
        pr.font.color.rgb = SAGE_80
        if purpose:
            pp = doc.add_paragraph()
            pp.paragraph_format.space_after = Pt(4)
            ppr = pp.add_run(purpose)
            ppr.font.name = BODY_FONT
            ppr.font.size = Pt(9.5)
            ppr.italic = True
            ppr.font.color.rgb = SAGE_60

        table = doc.add_table(rows=1, cols=3)
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        hdr = table.rows[0].cells
        _cell_text(hdr[0], "#", bold=True, color=WHITE, size=9.5,
                   font=HEADING_FONT, align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_text(hdr[1], "Step (5 words)", bold=True, color=WHITE, size=9.5,
                   font=HEADING_FONT)
        _cell_text(hdr[2], "What it means", bold=True, color=WHITE, size=9.5,
                   font=HEADING_FONT)
        for c in hdr:
            _shade(c, HEADER_BG)
            _set_borders(c, HEADER_BG)

        for i, step in enumerate(steps):
            sid = str(step.get("id") or "")
            label = str(step.get("label") or "").strip()
            expl = str(step.get("explanation") or "").strip()
            wc = _word_count(label)
            flagged = wc != 5
            if flagged:
                warnings.append(
                    f"{phase.get('id','?')}/{sid or i}: label has {wc} words (want 5): \"{label}\""
                )

            cells = table.add_row().cells
            _cell_text(cells[0], sid, color=SAGE_60, size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
            label_disp = label + ("  ⚠" if flagged else "")
            _cell_text(cells[1], label_disp, bold=True, color=SAGE_80, size=10.5)
            _cell_text(cells[2], expl, color=INK, size=10)

            bg = ZEBRA_BG if i % 2 == 0 else "FFFFFF"
            for c in cells:
                _shade(c, bg)
                _set_borders(c)

        _col_widths(table, col_w)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))

    return {
        "out": str(out_path),
        "phases": len(phases),
        "steps": total_steps,
        "label_warnings": warnings,
    }


def main():
    ap = argparse.ArgumentParser(description="Render a skill-breakdown JSON to a styled .docx")
    ap.add_argument("--in", dest="inp", help="input JSON file (default: stdin)")
    ap.add_argument("--out", dest="out", required=True, help="output .docx path")
    args = ap.parse_args()

    try:
        raw = Path(args.inp).read_text(encoding="utf-8") if args.inp else sys.stdin.read()
        data = json.loads(raw)
    except (OSError, json.JSONDecodeError) as e:
        sys.stderr.write(f"ERROR reading input JSON: {e}\n")
        sys.exit(2)

    try:
        result = render(data, Path(args.out))
    except (ValueError, KeyError) as e:
        sys.stderr.write(f"ERROR rendering: {e}\n")
        sys.exit(2)

    sys.stderr.write(
        f"OK  wrote {result['out']}  ({result['phases']} phases, {result['steps']} steps)\n"
    )
    if result["label_warnings"]:
        sys.stderr.write(
            f"WARN  {len(result['label_warnings'])} label(s) are not exactly 5 words:\n"
        )
        for w in result["label_warnings"]:
            sys.stderr.write(f"      - {w}\n")
    print(json.dumps(result, indent=2))


if __name__ == "__main__":
    main()
