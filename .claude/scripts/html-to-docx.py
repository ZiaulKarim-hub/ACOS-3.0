#!/usr/bin/env python3
"""
ACOS Loan Document Generator — HTML to Styled DOCX Converter

Two-pass approach:
  1. pandoc converts HTML → DOCX (reliable table/structure parsing)
  2. python-docx opens the result and applies institutional styling
     (navy headers, zebra tables, colored text, proper fonts)

Usage:
  python3 html-to-docx.py <input.html> <output.docx>
"""

import sys
import subprocess
import tempfile
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
except ImportError:
    print("ERROR: python-docx required. Install: pip install python-docx", file=sys.stderr)
    sys.exit(1)
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


# ── Style Constants ─────────────────────────────────────────────────────────

SAGE_80 = RGBColor(0x46, 0x5D, 0x53)     # OKOA primary — replaces navy
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
DARK_TEXT = RGBColor(0x16, 0x16, 0x16)    # OKOA ink
CORAL = RGBColor(0xFF, 0x79, 0x5E)        # OKOA accent
HEADER_BG = "465D53"                       # sage-80
ZEBRA_BG = "EEF3F0"                       # sage-05
FOOTER_BG = "DCE7E1"                      # sage-10

HEADING_FONT = "IBM Plex Sans"
BODY_FONT = "IBM Plex Sans"  # Must match pdf-styles.css primary: IBM Plex Sans
TABLE_FONT = "IBM Plex Sans"


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    # Remove existing shading
    for existing in tcPr.findall(f'{{{tcPr.nsmap.get("w", "")}}}shd'):
        tcPr.remove(existing)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    tcPr.append(shading)


def style_document(doc):
    """Apply institutional styling to an existing python-docx Document."""

    # ── Style headings ──────────────────────────────────────────────────
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            for run in para.runs:
                run.font.color.rgb = SAGE_80
                run.font.name = HEADING_FONT
            if para.style.name == "Heading 1":
                for run in para.runs:
                    run.font.size = Pt(18)
                    run.bold = True
            elif para.style.name == "Heading 2":
                for run in para.runs:
                    run.font.size = Pt(14)
                    run.bold = True
            elif para.style.name == "Heading 3":
                for run in para.runs:
                    run.font.size = Pt(11)
                    run.bold = True
        else:
            # Body text
            for run in para.runs:
                if run.font.name is None or run.font.name == "Times New Roman":
                    run.font.name = BODY_FONT
                if run.font.size is None:
                    run.font.size = Pt(10.5)
                if run.font.color.rgb is None:
                    run.font.color.rgb = DARK_TEXT

    # ── Style tables ────────────────────────────────────────────────────
    for table in doc.tables:
        rows = table.rows
        if not rows:
            continue

        num_cols = len(table.columns)

        for row_idx, row in enumerate(rows):
            is_header = row_idx == 0
            is_footer = row_idx == len(rows) - 1 and _is_footer_row(row)
            is_even = row_idx % 2 == 0 and not is_header

            for col_idx, cell in enumerate(row.cells):
                # Apply font styling to all cell content
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.name = TABLE_FONT
                        run.font.size = Pt(9)

                if is_header:
                    # Navy background, white bold text
                    set_cell_shading(cell, HEADER_BG)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.color.rgb = WHITE
                            run.bold = True
                            run.font.size = Pt(9)
                elif is_footer:
                    # Light navy background, bold text
                    set_cell_shading(cell, FOOTER_BG)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                            run.font.color.rgb = DARK_TEXT
                elif is_even and not is_header:
                    # Zebra striping
                    set_cell_shading(cell, ZEBRA_BG)
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.color.rgb = DARK_TEXT
                else:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.font.color.rgb = DARK_TEXT

    # ── Set margins ─────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.3)
        section.bottom_margin = Cm(2.8)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)


def _is_footer_row(row):
    """Heuristic: footer rows typically have bold text and words like 'Total' or 'Composite'."""
    text = ""
    for cell in row.cells:
        for para in cell.paragraphs:
            text += para.text.lower()
    return any(kw in text for kw in ["total", "composite", "overall", "score"])


def convert(html_path, docx_path):
    """Two-pass HTML → styled DOCX conversion."""

    # Pass 1: pandoc for structure
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
        tmp_path = tmp.name

    try:
        result = subprocess.run(
            ["pandoc", html_path, "-f", "html", "-o", tmp_path,
             "--wrap=none", "--standalone"],
            capture_output=True, text=True
        )

        if result.returncode != 0:
            print(f"ERROR: pandoc failed: {result.stderr}", file=sys.stderr)
            sys.exit(1)

        # Pass 2: python-docx for styling
        doc = Document(tmp_path)
        style_document(doc)
        doc.save(docx_path)
    finally:
        # Clean up temp file (also on pandoc failure / early exit)
        Path(tmp_path).unlink(missing_ok=True)

    print(f"DOCX generated: {docx_path}", file=sys.stderr)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("Usage: python3 html-to-docx.py <input.html> <output.docx>", file=sys.stderr)
        sys.exit(1)

    convert(sys.argv[1], sys.argv[2])
