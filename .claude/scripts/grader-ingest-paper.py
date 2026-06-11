#!/usr/bin/env python3
"""
grader-ingest-paper.py — Extract plain text from a single paper for the
acos-grader pipeline, including content from embedded images via Claude vision.

Handles:
- .txt              → copy verbatim
- .docx             → python-docx paragraph + table text + embedded-image vision OCR
- .pdf (text layer) → PyMuPDF text + per-page embedded-image vision OCR
- .pdf (scan)       → PyMuPDF renders each page, vision OCR (tesseract fallback)
- .png/.jpg/.jpeg/.tif/.tiff/.bmp/.webp → vision OCR (tesseract fallback)

Vision model: claude-sonnet-4-6 via anthropic SDK. Accepts either:
  - ANTHROPIC_API_KEY        → standard API account (x-api-key header)
  - ANTHROPIC_AUTH_TOKEN     → Claude subscription bearer token from
                               `claude setup-token` (Authorization: Bearer)
If both are set, ANTHROPIC_AUTH_TOKEN wins (subscription billing preferred).
Falls back to tesseract if: SDK missing, both env vars missing, API error, or --no-vision.

Vision prompt enforces structured output:
  <transcription>  verbatim student text / numbers / formulas  </transcription>
  <description>    structural description of diagrams / charts </description>
Only the transcription is treated as student content; description is prefixed
with a [FIGURE DESCRIPTION: ...] marker so graders can tell Claude's interpretation
apart from what the student wrote.

Usage:
    python3 grader-ingest-paper.py --input paper.pdf --output paper.txt
    python3 grader-ingest-paper.py --input scan.png --output scan.txt --no-vision
"""

from __future__ import annotations

import argparse
import base64
import json
import os
import shutil
import subprocess
import sys
import tempfile
import time
from pathlib import Path


# ---------- Configuration --------------------------------------------------


TEXT_LAYER_MIN_CHARS = 100
OCR_DPI = 300
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp", ".webp"}

VISION_MODEL_DEFAULT = "claude-sonnet-4-6"
VISION_MAX_TOKENS = 4096
VISION_TIMEOUT_SECONDS = 90
VISION_RETRIES = 1

DEFAULT_IMAGE_WARN_THRESHOLD = 200

VISION_PROMPT = """\
You are extracting the content of an image found inside a student's finance exam or
assignment submission. Output will be fed to a grader, so precision matters and any
commentary you add will be treated as the student's own writing unless properly fenced.

Produce TWO blocks, using these EXACT tags:

<transcription>
[Verbatim text, numbers, formulas, and cell values visible in the image. Include
handwritten content. Preserve line breaks. Mark illegible characters as [?].
Leave this block EMPTY if the image contains no readable text.]
</transcription>
<description>
[Structural description of any diagram, chart, flow, waterfall, tree, table, or
figure: axis labels, node labels, relationships, step sequences, directional
arrows. A grader reads this to understand what the student drew. Do NOT add
interpretation, correctness judgement, or commentary. Leave this block EMPTY
if the image is pure text.]
</description>

Context: {context_hint}

Return only the two blocks. No preamble. No closing remarks.
"""


# ---------- Extension routing ----------------------------------------------


def classify(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".txt":
        return "txt"
    if ext == ".docx":
        return "docx"
    if ext == ".pdf":
        return "pdf"
    if ext in IMAGE_EXTS:
        return "image"
    return "unknown"


# ---------- Vision (Claude) ------------------------------------------------


_anthropic_client = None
_vision_warned_missing = False

# Per-invocation counters for accurate mode reporting
_backend_stats = {"vision": 0, "tesseract": 0, "skipped": 0}


def _reset_backend_stats() -> None:
    _backend_stats["vision"] = 0
    _backend_stats["tesseract"] = 0
    _backend_stats["skipped"] = 0


def _backend_mode_label(use_vision: bool) -> str:
    v, t, s = _backend_stats["vision"], _backend_stats["tesseract"], _backend_stats["skipped"]
    if v == 0 and t == 0 and s == 0:
        return "no-images" if use_vision else "tesseract-only"
    if not use_vision:
        return "tesseract-only"
    if v > 0 and t == 0:
        return f"vision({v})"
    if v > 0 and t > 0:
        return f"vision({v})+tesseract-fallback({t})"
    return f"tesseract-fallback({t})"


def _get_anthropic_client():
    global _anthropic_client, _vision_warned_missing
    if _anthropic_client is not None:
        return _anthropic_client

    auth_token = os.environ.get("ANTHROPIC_AUTH_TOKEN")
    api_key = os.environ.get("ANTHROPIC_API_KEY")

    if not auth_token and not api_key:
        if not _vision_warned_missing:
            print(
                "WARNING: neither ANTHROPIC_AUTH_TOKEN nor ANTHROPIC_API_KEY is set; "
                "vision OCR disabled, falling back to tesseract for all images. "
                "Set one: run `claude setup-token` (subscription) or generate an API key.",
                file=sys.stderr,
            )
            _vision_warned_missing = True
        return None

    try:
        from anthropic import Anthropic
    except ImportError:
        if not _vision_warned_missing:
            print(
                "WARNING: anthropic SDK not installed; vision OCR disabled. "
                "Install via: pip install anthropic",
                file=sys.stderr,
            )
            _vision_warned_missing = True
        return None

    # Prefer subscription auth token when both are present (user has signalled
    # they want subscription billing over API billing)
    if auth_token:
        _anthropic_client = Anthropic(auth_token=auth_token)
    else:
        _anthropic_client = Anthropic(api_key=api_key)
    return _anthropic_client


def _vision_ocr(image_bytes: bytes, media_type: str, context_hint: str, model: str):
    """Call Claude vision on one image. Returns (transcription, description) or None on failure."""
    client = _get_anthropic_client()
    if client is None:
        return None

    b64 = base64.standard_b64encode(image_bytes).decode("ascii")
    prompt_text = VISION_PROMPT.format(context_hint=context_hint)

    for attempt in range(VISION_RETRIES + 1):
        try:
            response = client.messages.create(
                model=model,
                max_tokens=VISION_MAX_TOKENS,
                timeout=VISION_TIMEOUT_SECONDS,
                messages=[{
                    "role": "user",
                    "content": [
                        {
                            "type": "image",
                            "source": {
                                "type": "base64",
                                "media_type": media_type,
                                "data": b64,
                            },
                        },
                        {"type": "text", "text": prompt_text},
                    ],
                }],
            )
            raw = "".join(
                block.text for block in response.content
                if getattr(block, "type", None) == "text"
            )
            return _parse_vision_output(raw)
        except Exception as e:
            if attempt < VISION_RETRIES:
                time.sleep(1.5 ** attempt)
                continue
            print(
                f"WARNING: vision OCR failed ({type(e).__name__}: {str(e)[:200]}); "
                f"falling back to tesseract for this image",
                file=sys.stderr,
            )
            return None


def _parse_vision_output(raw: str):
    if not raw:
        return "", ""
    raw = raw.strip()
    # Strip markdown code fences if present
    if raw.startswith("```"):
        raw = raw.split("\n", 1)[1] if "\n" in raw else raw[3:]
        if raw.endswith("```"):
            raw = raw.rsplit("```", 1)[0]

    def _between(text: str, start_tag: str, end_tag: str) -> str:
        s = text.find(start_tag)
        if s == -1:
            return ""
        e = text.find(end_tag, s)
        if e == -1:
            return ""
        return text[s + len(start_tag):e].strip()

    transcription = _between(raw, "<transcription>", "</transcription>")
    description = _between(raw, "<description>", "</description>")

    # Defensive: if both tags missing, treat entire body as transcription
    if not transcription and not description:
        transcription = raw.strip()

    return transcription, description


# ---------- Image format normalization -------------------------------------


_DIRECT_MEDIA = {
    "image/jpeg": ("image/jpeg", "jpeg"),
    "image/jpg":  ("image/jpeg", "jpeg"),
    "image/png":  ("image/png",  "png"),
    "image/gif":  ("image/gif",  "gif"),
    "image/webp": ("image/webp", "webp"),
}

# PyMuPDF extract_image() 'ext' values → Anthropic media types
_PDF_EXT_MEDIA = {
    "jpeg": ("image/jpeg", "jpeg"),
    "jpg":  ("image/jpeg", "jpeg"),
    "png":  ("image/png",  "png"),
    "gif":  ("image/gif",  "gif"),
    "webp": ("image/webp", "webp"),
}


def _normalize_docx_image(content_type: str, blob: bytes):
    """Return (media_type, ext, bytes) suitable for Claude vision, or None."""
    ct = (content_type or "").strip().lower()
    if ct in _DIRECT_MEDIA:
        mt, ext = _DIRECT_MEDIA[ct]
        return mt, ext, blob
    # Convert TIFF, BMP, WMF, etc. to PNG via PyMuPDF
    try:
        import fitz
        with tempfile.NamedTemporaryFile(delete=False) as f:
            f.write(blob)
            tmp_path = f.name
        try:
            pix = fitz.Pixmap(tmp_path)
            if pix.n - pix.alpha >= 4:  # CMYK → RGB
                pix = fitz.Pixmap(fitz.csRGB, pix)
            return "image/png", "png", pix.tobytes("png")
        finally:
            Path(tmp_path).unlink(missing_ok=True)
    except Exception as e:
        print(f"WARNING: could not normalize image ({ct}): {e}", file=sys.stderr)
        return None


def _normalize_pdf_image(doc, xref: int, img_ext: str, img_bytes: bytes):
    """Return (media_type, ext, bytes) for a PDF-embedded image, or None."""
    key = (img_ext or "").lower()
    if key in _PDF_EXT_MEDIA:
        mt, ext = _PDF_EXT_MEDIA[key]
        return mt, ext, img_bytes
    # Exotic format (tiff, jpx, jb2, etc.) — rasterize to PNG
    try:
        import fitz
        pix = fitz.Pixmap(doc, xref)
        if pix.n - pix.alpha >= 4:
            pix = fitz.Pixmap(fitz.csRGB, pix)
        return "image/png", "png", pix.tobytes("png")
    except Exception as e:
        print(f"WARNING: could not rasterize PDF image xref={xref} ({img_ext}): {e}", file=sys.stderr)
        return None


# ---------- Tesseract (fallback) -------------------------------------------


def _tesseract_available() -> bool:
    return shutil.which("tesseract") is not None


def _tesseract_ocr_path(image_path: Path, lang: str = "eng") -> str:
    tesseract = shutil.which("tesseract")
    if not tesseract:
        return "[tesseract not installed]"
    try:
        result = subprocess.run(
            [tesseract, str(image_path), "-", "-l", lang, "--psm", "6"],
            capture_output=True, text=True, check=True, timeout=120,
        )
        return result.stdout
    except subprocess.CalledProcessError as e:
        return (e.stdout or "") + f"\n[OCR stderr: {e.stderr.strip()[:200]}]"
    except subprocess.TimeoutExpired:
        return "[OCR timed out after 120s]"


def _tesseract_ocr_bytes(image_bytes: bytes, ext: str, lang: str = "eng") -> str:
    with tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False) as tmp:
        tmp.write(image_bytes)
        tmp_path = Path(tmp.name)
    try:
        return _tesseract_ocr_path(tmp_path, lang=lang)
    finally:
        tmp_path.unlink(missing_ok=True)


# ---------- OCR dispatch ---------------------------------------------------


def _format_embedded(idx: int, location: str, transcription: str, description: str, backend: str) -> str:
    """Emit a text block for one embedded image.

    Transcription is printed plainly (treated as student content).
    Description, if any, is fenced with [FIGURE DESCRIPTION: ...] so graders
    can distinguish it from student writing.
    """
    tag_suffix = "" if backend == "vision" else f" ({backend})"
    header = f"[EMBEDDED IMAGE {idx}"
    if location:
        header += f" — {location}"
    header += f"]{tag_suffix}"

    parts: list[str] = [header]
    t = (transcription or "").strip()
    d = (description or "").strip()

    if t:
        parts.append(t)
        if d:
            parts.append(f"[FIGURE DESCRIPTION: {d}]")
    elif d:
        parts.append(f"[FIGURE DESCRIPTION: {d}]")
    else:
        parts.append("[no readable content]")

    return "\n".join(parts)


def _ocr_embedded_image(
    image_bytes: bytes,
    media_type: str,
    ext: str,
    idx: int,
    location: str,
    context_hint: str,
    use_vision: bool,
    lang: str,
    vision_model: str,
) -> str:
    """OCR one embedded image. Vision first if enabled; tesseract fallback."""
    if use_vision:
        result = _vision_ocr(image_bytes, media_type, context_hint, vision_model)
        if result is not None:
            transcription, description = result
            _backend_stats["vision"] += 1
            return _format_embedded(idx, location, transcription, description, backend="vision")
    # Tesseract fallback
    text = _tesseract_ocr_bytes(image_bytes, ext, lang=lang).strip()
    _backend_stats["tesseract"] += 1
    return _format_embedded(idx, location, text, "", backend="tesseract")


def _ocr_page_image(
    image_bytes: bytes,
    media_type: str,
    ext: str,
    page_num: int,
    context_hint: str,
    use_vision: bool,
    lang: str,
    vision_model: str,
) -> str:
    """OCR one full page of a scan/image-only PDF. No [EMBEDDED IMAGE] wrapper."""
    header = f"--- Page {page_num} ---"
    if use_vision:
        result = _vision_ocr(image_bytes, media_type, context_hint, vision_model)
        if result is not None:
            transcription, description = result
            _backend_stats["vision"] += 1
            lines: list[str] = [header]
            t = transcription.strip()
            d = description.strip()
            if t:
                lines.append(t)
            if d:
                lines.append(f"[FIGURE DESCRIPTION: {d}]")
            if len(lines) == 1:
                lines.append("[no readable content on this page]")
            return "\n".join(lines)
    # Tesseract fallback
    text = _tesseract_ocr_bytes(image_bytes, ext, lang=lang).strip()
    _backend_stats["tesseract"] += 1
    body = text if text else "[no readable content]"
    return f"{header}\n{body}"


# ---------- Per-format extractors ------------------------------------------


def extract_txt(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def extract_docx(
    path: Path,
    use_vision: bool,
    lang: str,
    image_warn_threshold: int,
    vision_model: str,
) -> str:
    try:
        from docx import Document
    except ImportError:
        sys.exit("ERROR: python-docx not installed. Run: pip install python-docx")

    doc = Document(path)
    parts: list[str] = []

    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    # Embedded images — walk ALL parts in the package (body + headers + footers
    # + footnotes + comments), not just the main document's direct rels.
    # Skip /docProps/* (thumbnails, metadata previews) — those are not student content.
    image_parts: list[tuple[str, bytes, str]] = []
    seen_partnames: set[str] = set()
    try:
        for part in doc.part.package.iter_parts():
            try:
                partname = str(part.partname)
                if partname in seen_partnames:
                    continue
                if partname.startswith("/docProps/"):
                    continue  # skip document-metadata thumbnails
                ct = (part.content_type or "").lower()
                if not ct.startswith("image/"):
                    continue
                blob = part.blob
                if not blob:
                    continue
                seen_partnames.add(partname)
                image_parts.append((ct, blob, partname))
            except Exception:
                continue
    except Exception as e:
        print(f"WARNING: could not enumerate DOCX image parts: {e}", file=sys.stderr)

    if image_parts:
        if len(image_parts) > image_warn_threshold:
            print(
                f"WARNING: {path.name} contains {len(image_parts)} embedded images "
                f"(> {image_warn_threshold}). Vision OCR cost may be significant.",
                file=sys.stderr,
            )
        parts.append("")  # blank separator before image block
        for idx, (ct, blob, partname) in enumerate(image_parts, start=1):
            norm = _normalize_docx_image(ct, blob)
            if norm is None:
                parts.append(f"[EMBEDDED IMAGE {idx}: unsupported content-type {ct}]")
                continue
            media_type, ext, bytes_to_send = norm
            location = f"DOCX part {partname}"
            context = (
                f"Embedded image #{idx} inside a DOCX submission ({path.name}). "
                f"This may contain handwritten work, a diagram, a chart, a screenshot of a spreadsheet, "
                f"or an equation image. Transcribe all text/numbers verbatim."
            )
            parts.append(
                _ocr_embedded_image(
                    bytes_to_send, media_type, ext, idx, location, context,
                    use_vision, lang, vision_model,
                )
            )

    return "\n".join(parts)


def extract_pdf(
    path: Path,
    use_vision: bool,
    lang: str,
    image_warn_threshold: int,
    vision_model: str,
) -> str:
    try:
        import fitz
    except ImportError:
        sys.exit("ERROR: PyMuPDF not installed. Run: pip install pymupdf")

    with fitz.open(path) as doc:
        text_layer = "\n".join((page.get_text() or "") for page in doc)

        if len(text_layer.strip()) >= TEXT_LAYER_MIN_CHARS:
            # Typed PDF: use text layer, also OCR any embedded images (charts,
            # screenshots, hand-drawn diagrams pasted into an otherwise-typed doc).
            parts: list[str] = [text_layer]

            # Collect unique images across all pages
            image_refs: list[tuple[int, int]] = []  # (page_num, xref)
            seen_xrefs: set[int] = set()
            for page_num, page in enumerate(doc, start=1):
                try:
                    for info in page.get_images(full=True):
                        xref = info[0]
                        if xref in seen_xrefs:
                            continue
                        seen_xrefs.add(xref)
                        image_refs.append((page_num, xref))
                except Exception as e:
                    print(f"WARNING: get_images failed on page {page_num}: {e}", file=sys.stderr)

            if image_refs:
                if len(image_refs) > image_warn_threshold:
                    print(
                        f"WARNING: {path.name} contains {len(image_refs)} unique embedded images "
                        f"(> {image_warn_threshold}). Vision OCR cost may be significant.",
                        file=sys.stderr,
                    )
                parts.append("")  # blank line before images block
                for idx, (page_num, xref) in enumerate(image_refs, start=1):
                    try:
                        img_dict = doc.extract_image(xref)
                    except Exception as e:
                        parts.append(f"[EMBEDDED IMAGE {idx} on page {page_num}: extract failed ({e})]")
                        continue
                    if not img_dict:
                        continue
                    img_ext = (img_dict.get("ext") or "").lower()
                    img_bytes = img_dict.get("image") or b""
                    if not img_bytes:
                        continue
                    norm = _normalize_pdf_image(doc, xref, img_ext, img_bytes)
                    if norm is None:
                        parts.append(f"[EMBEDDED IMAGE {idx} on page {page_num}: unsupported format {img_ext}]")
                        continue
                    media_type, ext, bytes_to_send = norm
                    location = f"page {page_num}"
                    context = (
                        f"Embedded image on page {page_num} of a typed-PDF submission "
                        f"({path.name}). The PDF also has a text layer, so this image likely "
                        f"contains diagrams, charts, spreadsheets, or hand-drawn work that the "
                        f"typed text cannot convey."
                    )
                    parts.append(
                        _ocr_embedded_image(
                            bytes_to_send, media_type, ext, idx, location, context,
                            use_vision, lang, vision_model,
                        )
                    )

            return "\n".join(parts)

        # No text layer: render each page at OCR_DPI and vision-OCR as pages
        if len(doc) > image_warn_threshold:
            print(
                f"WARNING: {path.name} is a scan with {len(doc)} pages "
                f"(> {image_warn_threshold}). Vision OCR cost may be significant.",
                file=sys.stderr,
            )
        parts: list[str] = []
        with tempfile.TemporaryDirectory(prefix="grader-ocr-") as tmpdir:
            tmp_path = Path(tmpdir)
            for i, page in enumerate(doc, start=1):
                pix = page.get_pixmap(dpi=OCR_DPI)
                img_path = tmp_path / f"page-{i:04d}.png"
                pix.save(str(img_path))
                img_bytes = img_path.read_bytes()
                context = (
                    f"Scanned page {i} of {len(doc)} from a handwritten or image-based "
                    f"submission ({path.name}). Transcribe all student writing, equations, "
                    f"and numerical work verbatim. Describe any diagrams structurally."
                )
                parts.append(
                    _ocr_page_image(
                        img_bytes, "image/png", "png", i, context,
                        use_vision, lang, vision_model,
                    )
                )
        return "\n\n".join(parts)


def extract_image(
    path: Path,
    use_vision: bool,
    lang: str,
    vision_model: str,
) -> str:
    """OCR a single image file."""
    img_bytes = path.read_bytes()
    ext = path.suffix.lower().lstrip(".")
    skip_vision = False  # set when exotic-format conversion fails (bytes are not a vision-safe format)
    # Map to media type; coerce exotic formats to PNG
    if ext in ("jpg", "jpeg"):
        media_type = "image/jpeg"
    elif ext == "png":
        media_type = "image/png"
    elif ext == "gif":
        media_type = "image/gif"
    elif ext == "webp":
        media_type = "image/webp"
    else:
        # tif/tiff/bmp → convert to PNG
        try:
            import fitz
            pix = fitz.Pixmap(str(path))
            if pix.n - pix.alpha >= 4:
                pix = fitz.Pixmap(fitz.csRGB, pix)
            img_bytes = pix.tobytes("png")
            media_type = "image/png"
            ext = "png"
        except Exception as e:
            print(f"WARNING: could not convert {path.name} to PNG: {e}", file=sys.stderr)
            # Conversion failed: img_bytes are still the original TIFF/BMP. Sending
            # those to the vision API mislabeled as image/png gets rejected and
            # wastes a round-trip, so skip vision and go straight to tesseract
            # (which reads the original bytes via the original ext).
            skip_vision = True

    context = (
        f"Full-image submission ({path.name}). This is likely a scanned or photographed "
        f"answer sheet. Transcribe all student writing, equations, numbers, and tables "
        f"verbatim. Describe any diagrams structurally."
    )

    if use_vision and not skip_vision:
        result = _vision_ocr(img_bytes, media_type, context, vision_model)
        if result is not None:
            transcription, description = result
            _backend_stats["vision"] += 1
            lines: list[str] = []
            t = transcription.strip()
            d = description.strip()
            if t:
                lines.append(t)
            if d:
                lines.append(f"[FIGURE DESCRIPTION: {d}]")
            return "\n".join(lines) if lines else ""
    # Tesseract fallback
    _backend_stats["tesseract"] += 1
    return _tesseract_ocr_bytes(img_bytes, ext, lang=lang)


# ---------- Orchestration --------------------------------------------------


def ingest(
    input_path: Path,
    output_path: Path,
    lang: str,
    use_vision: bool,
    vision_model: str,
    image_warn_threshold: int,
) -> dict:
    kind = classify(input_path)
    if kind == "unknown":
        sys.exit(f"ERROR: Unsupported file type: {input_path.suffix}")

    _reset_backend_stats()

    if kind == "txt":
        text = extract_txt(input_path)
    elif kind == "docx":
        text = extract_docx(input_path, use_vision, lang, image_warn_threshold, vision_model)
    elif kind == "pdf":
        text = extract_pdf(input_path, use_vision, lang, image_warn_threshold, vision_model)
    elif kind == "image":
        text = extract_image(input_path, use_vision, lang, vision_model)
    else:
        sys.exit(f"ERROR: Unhandled kind: {kind}")

    text = (text or "").strip()

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(text + "\n", encoding="utf-8")

    return {
        "kind": kind,
        "chars": len(text),
        "empty": len(text) == 0,
        "mode": _backend_mode_label(use_vision),
    }


# ---------- Task()-based orchestration modes ------------------------------
#
# When ANTHROPIC_AUTH_TOKEN / ANTHROPIC_API_KEY are unavailable (e.g. the user
# only has a Claude subscription, no API access), the SDK-based vision path
# is dead. In that case, the orchestrator calls this script in --plan-out
# mode to extract text + dump images to disk, then spawns the
# grader-vision-ocr agent via Task() for each image (inheriting session auth),
# then calls this script in --assemble-from mode to stitch the final text.


def plan(input_path: Path, plan_out: Path, staging_dir: Path, image_warn_threshold: int) -> dict:
    """Extract text parts + dump images to staging_dir. Write plan.json. No vision calls."""
    kind = classify(input_path)
    if kind == "unknown":
        sys.exit(f"ERROR: Unsupported file type: {input_path.suffix}")

    staging_dir.mkdir(parents=True, exist_ok=True)
    assembly: list[dict] = []

    if kind == "txt":
        text = extract_txt(input_path).strip()
        if text:
            assembly.append({"type": "text", "content": text})
    elif kind == "docx":
        assembly = _plan_docx(input_path, staging_dir, image_warn_threshold)
    elif kind == "pdf":
        assembly = _plan_pdf(input_path, staging_dir, image_warn_threshold)
    elif kind == "image":
        ext = input_path.suffix.lstrip(".").lower()
        dest = staging_dir / f"page-001.{ext or 'png'}"
        shutil.copy(input_path, dest)
        assembly.append({
            "type": "page_image",
            "idx": 1,
            "page_num": 1,
            "image_path": str(dest),
            "context": (
                f"Full-image submission ({input_path.name}). Likely a scanned or "
                f"photographed answer sheet. Transcribe all writing verbatim, "
                f"describe any diagrams structurally."
            ),
        })

    plan_doc = {
        "paper_id": input_path.stem,
        "source": str(input_path),
        "kind": kind,
        "staging_dir": str(staging_dir),
        "assembly": assembly,
    }
    plan_out.parent.mkdir(parents=True, exist_ok=True)
    plan_out.write_text(json.dumps(plan_doc, indent=2), encoding="utf-8")

    n_text = sum(1 for e in assembly if e["type"] == "text")
    n_imgs = sum(1 for e in assembly if e["type"] in ("embedded_image", "page_image"))
    return {"text_parts": n_text, "images": n_imgs}


def _plan_docx(path: Path, staging_dir: Path, image_warn_threshold: int) -> list[dict]:
    try:
        from docx import Document
    except ImportError:
        sys.exit("ERROR: python-docx not installed. Run: pip install python-docx")
    doc = Document(path)
    out: list[dict] = []

    for p in doc.paragraphs:
        if p.text.strip():
            out.append({"type": "text", "content": p.text})

    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                out.append({"type": "text", "content": " | ".join(cells)})

    image_parts: list[tuple[str, bytes, str]] = []
    seen_partnames: set[str] = set()
    try:
        for part in doc.part.package.iter_parts():
            try:
                partname = str(part.partname)
                if partname in seen_partnames:
                    continue
                if partname.startswith("/docProps/"):
                    continue
                ct = (part.content_type or "").lower()
                if not ct.startswith("image/"):
                    continue
                blob = part.blob
                if not blob:
                    continue
                seen_partnames.add(partname)
                image_parts.append((ct, blob, partname))
            except Exception:
                continue
    except Exception as e:
        print(f"WARNING: could not enumerate DOCX image parts: {e}", file=sys.stderr)

    if len(image_parts) > image_warn_threshold:
        print(
            f"WARNING: {path.name} contains {len(image_parts)} embedded images "
            f"(> {image_warn_threshold}). Vision cost may be significant.",
            file=sys.stderr,
        )

    for idx, (ct, blob, partname) in enumerate(image_parts, start=1):
        norm = _normalize_docx_image(ct, blob)
        if norm is None:
            out.append({
                "type": "text",
                "content": f"[EMBEDDED IMAGE {idx}: unsupported content-type {ct}]",
            })
            continue
        _media_type, ext, bytes_to_save = norm
        img_path = staging_dir / f"img-{idx:03d}.{ext}"
        img_path.write_bytes(bytes_to_save)
        out.append({
            "type": "embedded_image",
            "idx": idx,
            "image_path": str(img_path),
            "location": f"DOCX part {partname}",
            "context": (
                f"Embedded image #{idx} inside a DOCX submission ({path.name}). "
                f"May contain handwritten work, a diagram, a chart, a spreadsheet "
                f"screenshot, or an equation image. Transcribe all text/numbers verbatim."
            ),
        })

    return out


def _plan_pdf(path: Path, staging_dir: Path, image_warn_threshold: int) -> list[dict]:
    try:
        import fitz
    except ImportError:
        sys.exit("ERROR: PyMuPDF not installed. Run: pip install pymupdf")
    out: list[dict] = []

    with fitz.open(path) as doc:
        text_layer = "\n".join((page.get_text() or "") for page in doc)

        if len(text_layer.strip()) >= TEXT_LAYER_MIN_CHARS:
            out.append({"type": "text", "content": text_layer})

            image_refs: list[tuple[int, int]] = []
            seen_xrefs: set[int] = set()
            for page_num, page in enumerate(doc, start=1):
                try:
                    for info in page.get_images(full=True):
                        xref = info[0]
                        if xref in seen_xrefs:
                            continue
                        seen_xrefs.add(xref)
                        image_refs.append((page_num, xref))
                except Exception:
                    continue

            if len(image_refs) > image_warn_threshold:
                print(
                    f"WARNING: {path.name} contains {len(image_refs)} unique embedded "
                    f"images (> {image_warn_threshold}). Vision cost may be significant.",
                    file=sys.stderr,
                )

            for idx, (page_num, xref) in enumerate(image_refs, start=1):
                try:
                    img_dict = doc.extract_image(xref)
                except Exception as e:
                    out.append({
                        "type": "text",
                        "content": f"[EMBEDDED IMAGE {idx} on page {page_num}: extract failed ({e})]",
                    })
                    continue
                if not img_dict:
                    continue
                img_ext = (img_dict.get("ext") or "").lower()
                img_bytes = img_dict.get("image") or b""
                if not img_bytes:
                    continue
                norm = _normalize_pdf_image(doc, xref, img_ext, img_bytes)
                if norm is None:
                    out.append({
                        "type": "text",
                        "content": f"[EMBEDDED IMAGE {idx} on page {page_num}: unsupported format {img_ext}]",
                    })
                    continue
                _media_type, ext, bytes_to_save = norm
                img_path = staging_dir / f"img-{idx:03d}.{ext}"
                img_path.write_bytes(bytes_to_save)
                out.append({
                    "type": "embedded_image",
                    "idx": idx,
                    "image_path": str(img_path),
                    "location": f"page {page_num}",
                    "context": (
                        f"Embedded image on page {page_num} of a typed-PDF submission "
                        f"({path.name}). The PDF also has a text layer, so this image "
                        f"likely contains diagrams, charts, spreadsheets, or hand-drawn "
                        f"work that the typed text cannot convey."
                    ),
                })
        else:
            if len(doc) > image_warn_threshold:
                print(
                    f"WARNING: {path.name} is a scan with {len(doc)} pages "
                    f"(> {image_warn_threshold}). Vision cost may be significant.",
                    file=sys.stderr,
                )
            for i, page in enumerate(doc, start=1):
                pix = page.get_pixmap(dpi=OCR_DPI)
                img_path = staging_dir / f"page-{i:03d}.png"
                pix.save(str(img_path))
                out.append({
                    "type": "page_image",
                    "idx": i,
                    "page_num": i,
                    "image_path": str(img_path),
                    "context": (
                        f"Scanned page {i} of {len(doc)} from a handwritten or image-based "
                        f"submission ({path.name}). Transcribe all student writing, "
                        f"equations, and numerical work verbatim. Describe diagrams structurally."
                    ),
                })

    return out


def assemble(plan_path: Path, output_path: Path, vision_results_dir: Path) -> dict:
    """Stitch text parts + vision-produced image transcriptions into final paper text."""
    plan_doc = json.loads(plan_path.read_text(encoding="utf-8"))
    assembly = plan_doc.get("assembly", [])

    parts: list[str] = []
    resolved = 0
    missing = 0

    for entry in assembly:
        t = entry.get("type")
        if t == "text":
            parts.append(entry.get("content", ""))
        elif t == "embedded_image":
            idx = int(entry.get("idx", 0))
            location = entry.get("location", "")
            vision_text = _load_vision_result(vision_results_dir, idx, page_style=False)
            if vision_text is None:
                missing += 1
                parts.append(_format_embedded(idx, location, "", "[vision result missing]", backend="missing"))
            else:
                transcription, description = _parse_vision_output(vision_text)
                resolved += 1
                parts.append(_format_embedded(idx, location, transcription, description, backend="vision-task"))
        elif t == "page_image":
            idx = int(entry.get("idx", 0))
            page_num = int(entry.get("page_num", idx))
            header = f"--- Page {page_num} ---"
            vision_text = _load_vision_result(vision_results_dir, idx, page_style=True)
            if vision_text is None:
                missing += 1
                parts.append(f"{header}\n[vision result missing]")
            else:
                transcription, description = _parse_vision_output(vision_text)
                resolved += 1
                lines = [header]
                if transcription.strip():
                    lines.append(transcription.strip())
                if description.strip():
                    lines.append(f"[FIGURE DESCRIPTION: {description.strip()}]")
                if len(lines) == 1:
                    lines.append("[no readable content on this page]")
                parts.append("\n".join(lines))

    text = "\n".join(parts).strip()
    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(text + "\n", encoding="utf-8")

    return {
        "chars": len(text),
        "empty": len(text) == 0,
        "resolved": resolved,
        "missing": missing,
    }


def _load_vision_result(results_dir: Path, idx: int, page_style: bool) -> str | None:
    """Look up a vision result file for image idx.

    Agents write to results_dir/img-<idx:03d>.txt (embedded) or
    results_dir/page-<idx:03d>.txt (scan pages). Both prefixes are probed for
    flexibility; page_style selects the preferred prefix order.
    """
    prefixes = ("page", "img") if page_style else ("img", "page")
    for prefix in prefixes:
        p = results_dir / f"{prefix}-{idx:03d}.txt"
        if p.exists():
            return p.read_text(encoding="utf-8")
    return None


# ---------- Main -----------------------------------------------------------


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--input", help="Path to input paper (txt/docx/pdf/image). Required for direct/plan modes.")
    ap.add_argument("--output", help="Path to write extracted text. Required for direct/assemble modes.")
    ap.add_argument("--lang", default="eng", help="Tesseract language code (default: eng)")
    ap.add_argument(
        "--no-vision", action="store_true",
        help="Disable Claude vision OCR; use tesseract only (old behavior)",
    )
    ap.add_argument(
        "--vision-model", default=VISION_MODEL_DEFAULT,
        help=f"Claude vision model (default: {VISION_MODEL_DEFAULT})",
    )
    ap.add_argument(
        "--max-images", type=int, default=DEFAULT_IMAGE_WARN_THRESHOLD,
        help=(
            f"Warn (to stderr) if an input contains more than this many embedded images "
            f"(default: {DEFAULT_IMAGE_WARN_THRESHOLD}). Does not abort."
        ),
    )

    # Task()-based orchestration modes
    ap.add_argument(
        "--plan-out",
        help=(
            "Extract text + dump images to --staging-dir, write a plan.json at this "
            "path, then exit. No vision OCR calls. Used by the Task()-based "
            "orchestrator when no Anthropic API credentials are available."
        ),
    )
    ap.add_argument(
        "--staging-dir",
        help="Directory to dump embedded images to (required with --plan-out).",
    )
    ap.add_argument(
        "--assemble-from",
        help=(
            "Path to a plan.json previously produced by --plan-out. Reads the plan "
            "and stitches text parts together with vision results from "
            "--vision-results-dir, writing the final paper text to --output."
        ),
    )
    ap.add_argument(
        "--vision-results-dir",
        help=(
            "Directory containing per-image vision outputs named img-NNN.txt / "
            "page-NNN.txt (required with --assemble-from). Each file contains "
            "<transcription>...</transcription><description>...</description>."
        ),
    )

    args = ap.parse_args()

    # Mode dispatch: plan > assemble > direct
    if args.plan_out:
        if not args.input:
            sys.exit("ERROR: --plan-out requires --input")
        if not args.staging_dir:
            sys.exit("ERROR: --plan-out requires --staging-dir")
        input_path = Path(args.input).expanduser().resolve()
        plan_out = Path(args.plan_out).expanduser().resolve()
        staging_dir = Path(args.staging_dir).expanduser().resolve()
        if not input_path.exists() or not input_path.is_file():
            sys.exit(f"ERROR: Input file not found: {input_path}")
        result = plan(input_path, plan_out, staging_dir, args.max_images)
        print(
            f"PLANNED paper_id={input_path.stem} text_parts={result['text_parts']} "
            f"images={result['images']} → {plan_out}"
        )
        return 0

    if args.assemble_from:
        if not args.output:
            sys.exit("ERROR: --assemble-from requires --output")
        if not args.vision_results_dir:
            sys.exit("ERROR: --assemble-from requires --vision-results-dir")
        plan_path = Path(args.assemble_from).expanduser().resolve()
        output_path = Path(args.output).expanduser().resolve()
        vision_dir = Path(args.vision_results_dir).expanduser().resolve()
        if not plan_path.exists():
            sys.exit(f"ERROR: plan file not found: {plan_path}")
        result = assemble(plan_path, output_path, vision_dir)
        status = "EMPTY" if result["empty"] else "OK"
        print(
            f"ASSEMBLED chars={result['chars']} resolved={result['resolved']} "
            f"missing={result['missing']} status={status} → {output_path}"
        )
        return 2 if result["empty"] else 0

    # Direct mode (original behavior: one-shot extract + vision)
    if not args.input or not args.output:
        sys.exit("ERROR: --input and --output are required for direct mode")

    input_path = Path(args.input).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()

    if not input_path.exists() or not input_path.is_file():
        sys.exit(f"ERROR: Input file not found: {input_path}")

    use_vision = not args.no_vision

    result = ingest(
        input_path=input_path,
        output_path=output_path,
        lang=args.lang,
        use_vision=use_vision,
        vision_model=args.vision_model,
        image_warn_threshold=args.max_images,
    )

    status = "EMPTY" if result["empty"] else "OK"
    print(
        f"INGESTED kind={result['kind']} chars={result['chars']} "
        f"mode={result['mode']} status={status} → {output_path}"
    )

    if result["empty"]:
        return 2
    return 0


if __name__ == "__main__":
    sys.exit(main())
