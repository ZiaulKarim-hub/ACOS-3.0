#!/usr/bin/env python3
"""
grader-parse-rubric.py — Parse DOCX / PDF / XLSX rubric into acos-grader
internal YAML format.

Input: a rubric file (DOCX, PDF, or XLSX) with criteria, points, descriptions.
Output: YAML conforming to `.claude/skills/acos-grader/templates/rubric-schema.yaml`.

For best results the rubric should have table-like structure with columns for
criterion name, points, and description. Free-form prose rubrics are best
pre-processed via `acos-data-extractor` before invoking this script.

Usage:
    python3 grader-parse-rubric.py --input rubric.docx --output rubric.yaml
    python3 grader-parse-rubric.py --input rubric.xlsx --output rubric.yaml --total-points 100
"""

from __future__ import annotations

import argparse
import re
import sys
from dataclasses import dataclass, asdict
from datetime import datetime, timezone
from pathlib import Path
from typing import Iterable


# ---------- Data model ------------------------------------------------------


@dataclass
class Criterion:
    id: str
    name: str
    points: float
    description: str


# ---------- Utility helpers -------------------------------------------------


def _slugify(name: str, index: int) -> str:
    slug = re.sub(r"[^a-z0-9]+", "_", name.lower()).strip("_")
    if not slug:
        slug = f"criterion_{index + 1}"
    # Cap length
    return slug[:40] if len(slug) > 40 else slug


def _normalize_points(raw: str) -> float | None:
    """Parse '10', '10 pts', '10 points', '10.5' → float. Return None if not a number."""
    if raw is None:
        return None
    s = str(raw).strip().lower().replace(",", "")
    s = re.sub(r"(pts?|points?|marks?)\b", "", s).strip()
    try:
        return float(s)
    except ValueError:
        return None


# ---------- Extractors by format -------------------------------------------


def extract_docx(path: Path) -> list[Criterion]:
    try:
        from docx import Document  # python-docx
    except ImportError:
        sys.exit("ERROR: python-docx not installed. Run: pip install python-docx")

    doc = Document(path)
    criteria: list[Criterion] = []

    # Prefer tables — rubrics are typically tabular
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        if len(rows) < 2:
            continue
        header = [h.lower() for h in rows[0]]
        if not _header_looks_like_rubric(header):
            continue

        # Identify columns
        name_col = _find_col(header, ("criterion", "criteria", "item", "question"))
        points_col = _find_col(header, POINTS_KEYS)
        desc_col = _find_col(header, ("description", "details", "rubric", "notes"))

        if name_col is None or points_col is None:
            continue

        for i, row in enumerate(rows[1:]):
            if len(row) <= max(name_col, points_col):
                continue
            name = row[name_col].strip()
            pts = _normalize_points(row[points_col])
            desc = row[desc_col].strip() if desc_col is not None and desc_col < len(row) else ""
            if not name or pts is None or pts <= 0:
                continue
            criteria.append(Criterion(
                id=_slugify(name, i),
                name=name,
                points=pts,
                description=desc or name,
            ))

    if criteria:
        return criteria

    # Fallback: parse numbered list in paragraphs
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    return _parse_numbered_list(paragraphs)


def extract_xlsx(path: Path) -> list[Criterion]:
    try:
        from openpyxl import load_workbook
    except ImportError:
        sys.exit("ERROR: openpyxl not installed. Run: pip install openpyxl")

    wb = load_workbook(path, data_only=True)
    criteria: list[Criterion] = []

    for sheet in wb.worksheets:
        rows = [[str(c.value).strip() if c.value is not None else "" for c in row] for row in sheet.iter_rows()]
        if len(rows) < 2:
            continue
        header = [h.lower() for h in rows[0]]
        if not _header_looks_like_rubric(header):
            continue

        name_col = _find_col(header, ("criterion", "criteria", "item", "question"))
        points_col = _find_col(header, POINTS_KEYS)
        desc_col = _find_col(header, ("description", "details", "rubric", "notes"))

        if name_col is None or points_col is None:
            continue

        for i, row in enumerate(rows[1:]):
            if len(row) <= max(name_col, points_col):
                continue
            name = row[name_col].strip()
            pts = _normalize_points(row[points_col])
            desc = row[desc_col].strip() if desc_col is not None and desc_col < len(row) else ""
            if not name or pts is None or pts <= 0:
                continue
            criteria.append(Criterion(
                id=_slugify(name, i),
                name=name,
                points=pts,
                description=desc or name,
            ))
        if criteria:
            break  # First matching sheet wins

    return criteria


def extract_pdf(path: Path) -> list[Criterion]:
    try:
        import pdfplumber
    except ImportError:
        sys.exit("ERROR: pdfplumber not installed. Run: pip install pdfplumber")

    criteria: list[Criterion] = []

    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            for table in page.extract_tables() or []:
                if len(table) < 2:
                    continue
                header = [(h or "").lower().strip() for h in table[0]]
                if not _header_looks_like_rubric(header):
                    continue

                name_col = _find_col(header, ("criterion", "criteria", "item", "question"))
                points_col = _find_col(header, POINTS_KEYS)
                desc_col = _find_col(header, ("description", "details", "rubric", "notes"))

                if name_col is None or points_col is None:
                    continue

                for i, row in enumerate(table[1:]):
                    if row is None or len(row) <= max(name_col, points_col):
                        continue
                    name = (row[name_col] or "").strip()
                    pts = _normalize_points(row[points_col] or "")
                    desc = (row[desc_col] or "").strip() if desc_col is not None and desc_col < len(row) else ""
                    if not name or pts is None or pts <= 0:
                        continue
                    criteria.append(Criterion(
                        id=_slugify(name, i),
                        name=name,
                        points=pts,
                        description=desc or name,
                    ))
            if criteria:
                break
        # If no table, fall back to numbered list in all-text extraction
        if not criteria:
            text = "\n".join((page.extract_text() or "") for page in pdf.pages)
            paragraphs = [p.strip() for p in text.split("\n") if p.strip()]
            criteria = _parse_numbered_list(paragraphs)

    return criteria


# ---------- Heuristics -----------------------------------------------------


# Shared keyword set for the points column. Includes the SINGULAR 'point' so a
# header literally labeled 'Point' both passes _header_looks_like_rubric AND
# resolves a points column in _find_col (previously asymmetric — the gate
# accepted 'point' but _find_col did not, silently dropping the table).
POINTS_KEYS = ("point", "points", "pts", "marks", "score")


def _header_looks_like_rubric(header: list[str]) -> bool:
    joined = " ".join(header)
    has_criterion = any(k in joined for k in ("criterion", "criteria", "item", "question"))
    has_points = any(k in joined for k in POINTS_KEYS)
    return has_criterion and has_points


def _find_col(header: list[str], keys: Iterable[str]) -> int | None:
    for i, h in enumerate(header):
        for k in keys:
            if k in h:
                return i
    return None


def _parse_numbered_list(paragraphs: list[str]) -> list[Criterion]:
    """Fallback parser: matches lines like '1. Name ... (10 points)'"""
    criteria: list[Criterion] = []
    pattern = re.compile(
        r"^\s*(?:\d+[\.\)]|\-|\*)\s*(.+?)\s*[\(\[]?\s*(\d+(?:\.\d+)?)\s*(?:pts?|points?|marks?)[\)\]]?\s*$",
        re.IGNORECASE,
    )
    for i, p in enumerate(paragraphs):
        m = pattern.match(p)
        if not m:
            continue
        name = m.group(1).strip(" -—:").strip()
        pts = float(m.group(2))
        if name and pts > 0:
            criteria.append(Criterion(
                id=_slugify(name, i),
                name=name,
                points=pts,
                description=name,
            ))
    return criteria


# ---------- Orchestration --------------------------------------------------


def parse_rubric(input_path: Path) -> list[Criterion]:
    ext = input_path.suffix.lower()
    if ext == ".docx":
        return extract_docx(input_path)
    if ext == ".xlsx":
        return extract_xlsx(input_path)
    if ext == ".pdf":
        return extract_pdf(input_path)
    sys.exit(f"ERROR: Unsupported extension {ext}. Use .docx, .xlsx, or .pdf.")


def build_yaml(
    criteria: list[Criterion],
    source_path: Path,
    total_points_declared: float | None,
) -> str:
    total_computed = sum(c.points for c in criteria)

    if total_points_declared is not None and abs(total_computed - total_points_declared) > 0.5:
        sys.exit(
            f"ERROR: Criterion points sum to {total_computed} but declared total is "
            f"{total_points_declared}. Fix the rubric or the declared total."
        )

    total_points = total_points_declared if total_points_declared is not None else total_computed

    # Ensure unique IDs
    seen: dict[str, int] = {}
    for c in criteria:
        if c.id in seen:
            seen[c.id] += 1
            c.id = f"{c.id}_{seen[c.id]}"
        else:
            seen[c.id] = 1

    # Build YAML manually (avoid yaml dependency — stdlib only per ACOS convention)
    lines: list[str] = []
    lines.append("rubric:")
    lines.append(f"  title: {_yaml_str(source_path.stem)}")
    lines.append(f"  total_points: {total_points}")
    lines.append("  criteria:")
    for c in criteria:
        lines.append(f"    - id: {c.id}")
        lines.append(f"      name: {_yaml_str(c.name)}")
        lines.append(f"      points: {c.points}")
        lines.append(f"      description: {_yaml_str(c.description)}")
    lines.append("  metadata:")
    lines.append(f"    source_file: {_yaml_str(source_path.name)}")
    lines.append(f"    source_extension: {_yaml_str(source_path.suffix.lstrip('.'))}")
    lines.append(f"    parsed_at: {_yaml_str(datetime.now(timezone.utc).isoformat())}")
    lines.append("    parser: grader-parse-rubric.py")
    lines.append("    parser_version: '1.0'")
    return "\n".join(lines) + "\n"


def _yaml_str(s: str) -> str:
    """Emit a YAML-safe string. Quote if it contains special chars; otherwise bare.

    Double-quoted YAML scalars MUST have control characters (\\n, \\r, \\t, etc.)
    replaced with their escape sequences — raw newlines inside a quoted scalar
    split the value into adjacent lines, corrupting the document.
    """
    if not s:
        return '""'
    needs_quote = any(ch in s for ch in ":#{}[]|>'\"&*!%@`,\n\r\t\f\v")
    # YAML block indicators are only special in LEADING position. A scalar whose
    # value begins with one (e.g. "- Introduction to risk" reads as a sequence
    # entry; "? maybe later" as a complex mapping key) must be quoted or it fails
    # to round-trip through yaml.safe_load downstream.
    leading_indicators = set("-?:,[]{}#&*!|>%@`\"' \t")
    leading_special = s[:1] in leading_indicators or s[:2] in ("- ", "? ", ": ")
    if (
        needs_quote
        or leading_special
        or s.strip() != s
        or s.lower() in ("yes", "no", "true", "false", "null")
    ):
        escaped = (
            s.replace("\\", "\\\\")
             .replace('"', '\\"')
             .replace("\n", "\\n")
             .replace("\r", "\\r")
             .replace("\t", "\\t")
             .replace("\f", "\\f")
             .replace("\v", "\\v")
        )
        return f'"{escaped}"'
    return s


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__, formatter_class=argparse.RawDescriptionHelpFormatter)
    ap.add_argument("--input", required=True, help="Path to rubric file (.docx, .pdf, .xlsx)")
    ap.add_argument("--output", required=True, help="Path to write internal YAML rubric")
    ap.add_argument("--total-points", type=float, default=None,
                    help="Declared total points (optional; used to validate extracted sum)")
    args = ap.parse_args()

    input_path = Path(args.input).expanduser().resolve()
    output_path = Path(args.output).expanduser().resolve()

    if not input_path.exists():
        sys.exit(f"ERROR: Input file not found: {input_path}")
    if not input_path.is_file():
        sys.exit(f"ERROR: Input path is not a file: {input_path}")

    criteria = parse_rubric(input_path)

    if not criteria:
        sys.exit(
            f"ERROR: Could not extract any criteria from {input_path}. "
            f"Consider pre-processing via /acos-data-extractor or /acos-pdf-xlsx-converter."
        )

    yaml_text = build_yaml(criteria, input_path, args.total_points)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    output_path.write_text(yaml_text, encoding="utf-8")

    total_pts = sum(c.points for c in criteria)
    print(f"Parsed {len(criteria)} criteria, {total_pts} total points → {output_path}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
